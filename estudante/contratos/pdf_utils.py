from docx import Document
from io import BytesIO
from xhtml2pdf import pisa
from estudante.contratos.templates.field_list import FIELDS 
from django.db import connection

def consulta_matricula(matricula):
    with connection.cursor() as cursor:
        query = """
                    select
                        ee.nome_fantasia as "[NOME_EMPRESA]",
                        pde.nro_documento as "[CNPJ_EMPRESA]",
                        pee.logradouro || ' nº ' || pee.numero || ' na cidade de ' || pee.cidade || '/' || pee.estado as "[ENDERECO_EMPRESA]",
                        pp.nome as "[NOME_APRENDIZ]",
                        pe.logradouro || ' nº ' || pe.numero || ' na cidade de ' || pe.cidade || '/' || pe.estado as "[ENDERECO_APRENDIZ]",
                        pd.nro_documento as "[CPF_APRENDIZ]",
                        ec.descricao as "[OCUPACAO_APRENDIZ]",
                        ec.codigo as "[NRO_CBO]",
                        ec2.nome as "[NOME_CURSO]",
                        ec2.codigo as "[NRO_CURSO]",
                        '--' as "[PROTOCOLO_CURSO]",
                        em.quantidade_meses_contrato as "[QTD_MESES_CONTRATO]",
                        '--' as "[CBOS_ASSOCIADOS]",
                        TO_CHAR(em.data_inicio_contrato,
                        'DD/MM/YYYY') as "[INICIO_CONTRATO]",
                        TO_CHAR(em.data_terminio_contrato,
                        'DD/MM/YYYY') as "[TERMINO_CONTRATO]",
                        TO_CHAR(eat.data_inicio,
                        'DD/MM/YYYY') || ' a ' || TO_CHAR(eat.data_fim,
                        'DD/MM/YYYY') || ' das ' || eat.hora_inicio || ' às ' || eat.hora_termino as "[PERIODO_ATIVIDADE_TEORICA]",
                        --30/07/2024 a 15/07/2024, das 08:30 às 14:45
                                            (
                        select
                            (
                            select
                                STRING_AGG(ed.dia,
                                ', ')
                            from
                                estudante_turma et
                            inner join estudante_turma_dias_da_semana_empresa etddse on
                                et.id = etddse.turma_id
                            inner join estudante_diasemana ed on
                                ed.id = etddse.diasemana_id
                            where
                                et.id = em.turma_id) as dias_agregados
                        from
                            estudante_matricula em2
                        where
                            em2.id = em.id
                        limit 1) as "[DIAS_EMPRESA]",
                        em.hora_inicio_expediente as "[HORA_INICIO_EMPRESA]",
                        em.hora_fim_expediente as "[HORA_TERMINO_EMPRESA]",
                        (
                        select
                            (
                            select
                                STRING_AGG(ed.dia,
                                ', ')
                            from
                                estudante_turma et
                            inner join estudante_turma_dias_da_semana_curso etddsc on
                                et.id = etddsc.turma_id
                            inner join estudante_diasemana ed on
                                ed.id = etddsc.diasemana_id
                            where
                                et.id = em.turma_id) as dias_agregados
                        from
                            estudante_matricula em2
                        where
                            em2.id = em.id
                        limit 1) as "[DIAS_APRENDIZAGEM]",
                        em.salario as "[SALARIO]",
                        '--' as "[ATIVIDADES_PRATICAS]",
                        TO_CHAR(current_date,
                        'DD/MM/YYYY') as "[DATA_ATUAL]",
                        (
                        select
                            nome
                        from
                            pessoa_pessoa
                        where
                            id = pessoa_responsavel_id) as "[NOME_RESPONSAVEL_EMPRESA]",
                        (
                        select
                            nro_documento
                        from
                            pessoa_documento pd2
                        inner join pessoa_pessoa_documento ppd2 on
                            pd2.id = ppd2.documento_id
                            and ppd2.pessoa_id = pessoa_responsavel_id
                            and pd2.tipo_documento = 'CPF') as "[CPF_RESPONSAVEL_EMPRESA]"
                    from
                        estudante_matricula em,
                        pessoa_pessoa pp,
                        pessoa_endereco pe,
                        pessoa_endereco pee,
                        pessoa_pessoa_documento ppd,
                        estudante_empresa_documento eed,
                        pessoa_documento pd,
                        pessoa_documento pde,
                        estudante_turma et,
                        estudante_cbo ec,
                        estudante_curso ec2,
                        estudante_empresa ee,
                        estudante_atividade_teorica eat
                    where
                        em.pessoa_id = pp.id
                        and pp.endereco_id = pe.id
                        and ee.endereco_id = pee.id
                        and pp.id = ppd.pessoa_id
                        and ppd.documento_id = pd.id
                        and pd.tipo_documento = 'CPF'
                        and ee.id = eed.empresa_id
                        and eed.documento_id = pde.id
                        and pde.tipo_documento = 'CNPJ'
                        and em.turma_id = et.id
                        and em.cbo_id = ec.id
                        and em.curso_id = ec2.id
                        and em.empresa_id = ee.id
                        and em.id = eat.matricula_id
                        and em.numero_matricula = %s;
        """
        cursor.execute(query, [matricula])
        print('Pesquisa da matricula: ', matricula)
        resultados = cursor.fetchall()
        print('Resultado da query: ', resultados)
        colunas = [desc[0] for desc in cursor.description]
    
    dados = [dict(zip(colunas, resultado)) for resultado in resultados]
    return dados


def modify_docx(doc_path, matricula):
    dados = consulta_matricula(matricula)
    doc = Document(doc_path)
    
    for dado in dados:
        print('Resultado da pesquisa: ', dado)
        for search_text in FIELDS:
            print(f'Procurando: {search_text} para substituir por {dado[search_text]}')
            

            for para in doc.paragraphs:
                if search_text in para.text:
                    inline = para.runs
                    for i in range(len(inline)):
                        if search_text in inline[i].text:
                            inline[i].text = inline[i].text.replace(search_text, str(dado[search_text]))
            

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if search_text in cell.text:
                            inline = cell.paragraphs
                            for para in inline:
                                if search_text in para.text:
                                    for run in para.runs:
                                        if search_text in run.text:
                                            run.text = run.text.replace(search_text, str(dado[search_text]))
    

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
